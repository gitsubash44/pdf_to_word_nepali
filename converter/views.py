import os
import re
import tempfile
from django.shortcuts import render, redirect
from django.conf import settings
from django.http import FileResponse
from .models import PDFFile, ConvertedFile
from pdf2docx import Converter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import unicodedata


def home(request):
    return render(request, 'converter/convert.html')


def sanitize_filename(filename):
    # Normalize Unicode characters
    filename = unicodedata.normalize('NFKD', filename).encode('ascii', 'ignore').decode('ascii')
    # Remove special characters
    filename = re.sub(r'[^\w\s-]', '', filename)
    return filename

def convert_pdf_to_word(request):
    if request.method == 'POST':
        if not request.FILES.get('pdf_file'):
            return render(request, 'converter/convert.html', {
                'error': 'Please select a PDF file to upload'
            })
        
        try:
            pdf_file = request.FILES['pdf_file']
            
            # Sanitize the filename
            original_filename = pdf_file.name
            safe_filename = sanitize_filename(original_filename)
            
            # Save uploaded PDF
            pdf_model = PDFFile(file=pdf_file)
            pdf_model.save()
            
            # Create temp files
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                for chunk in pdf_file.chunks():
                    tmp_pdf.write(chunk)
                pdf_path = tmp_pdf.name
            
            # Ensure word_files directory exists
            word_dir = os.path.join(settings.MEDIA_ROOT, 'word_files')
            os.makedirs(word_dir, exist_ok=True)
            
            # Create output filename
            output_filename = f'converted_{safe_filename.replace(".pdf", ".docx")}'
            word_path = os.path.join(word_dir, output_filename)
            
            # Convert PDF to Word with enhanced Nepali font handling
            convert_pdf_with_nepali_fonts(pdf_path, word_path)
            
            # Clean up temp file
            os.unlink(pdf_path)
            
            # Save converted file to model
            converted_model = ConvertedFile(
                original_pdf=pdf_model,
                word_file=f'word_files/{output_filename}'
            )
            converted_model.save()
            
            return render(request, 'converter/result.html', {
                'original_file': pdf_model,
                'converted_file': converted_model
            })
            
        except Exception as e:
            return render(request, 'converter/convert.html', {
                'error': f'Conversion failed: {str(e)}'
            })
    
    return redirect('home')

def convert_pdf_with_nepali_fonts(pdf_path, word_path):
    """Enhanced conversion with Nepali font support"""
    # First convert using pdf2docx for layout
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()
    
    # Now enhance the document with proper font settings
    doc = Document(word_path)
    
    # Set Nepali font for the entire document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Try Preeti first, fallback to Arial Unicode MS
            run.font.name = 'Preeti'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Preeti')
            run.font.size = Pt(10)
    
    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Preeti'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Preeti')
                        run.font.size = Pt(10)
    
    doc.save(word_path)

def download_word_file(request, file_id):
    converted_file = ConvertedFile.objects.get(id=file_id)
    file_path = converted_file.word_file.path
    response = FileResponse(open(file_path, 'rb'))
    response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
    return response